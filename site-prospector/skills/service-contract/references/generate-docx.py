#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generates the LOCKED contract .docx (read-only, with editable regions for the client).
Usage: python3 generate-docx.py data.json output.docx
data.json: same keys as contract-template.html (CLIENT_NAME, VALUE, ... ) +
  "MAINTENANCE": true/false and "MAINTENANCE_VALUE" when hired.
Fields the CLIENT fills in (editable regions): CLIENT_DOC, CLIENT_ADDRESS (if they come
as "(to fill in)"), date and the client's signature."""
import json, sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

d = json.load(open(sys.argv[1], encoding='utf-8'))
PID = [100]

def par(doc, text='', bold=False, center=False, size=11, before=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(6)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Georgia'
    return p

def run(p, text, bold=False, size=11):
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Georgia'
    return r

def editable(p, text):
    """Inserts a snippet the client CAN edit (permStart/permEnd, everyone group)."""
    PID[0] += 1; pid = str(PID[0])
    ps = OxmlElement('w:permStart'); ps.set(qn('w:id'), pid); ps.set(qn('w:edGrp'), 'everyone')
    p._p.append(ps)
    r = run(p, text); r.font.highlight_color = 7  # yellow: shows where to fill in
    pe = OxmlElement('w:permEnd'); pe.set(qn('w:id'), pid)
    p._p.append(pe)

def field(p, value, label):
    """If the value came as '(to fill in)', it becomes an editable region; otherwise plain text."""
    if 'to fill' in (value or '').lower() or not value:
        editable(p, ' [' + label + ': fill in here] ')
    else:
        run(p, value)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.2); s.left_margin = s.right_margin = Cm(2.2)

par(doc, 'SERVICE CONTRACT', bold=True, center=True, size=13)
par(doc, 'CREATION AND PUBLISHING OF A WEB PAGE', bold=True, center=True, size=11)

p = par(doc); run(p, 'CLIENT: ', bold=True); run(p, d['CLIENT_NAME'] + ', ' + d['CLIENT_DOC_LABEL'] + ' no. ')
field(p, d.get('CLIENT_DOC'), 'CPF/CNPJ'); run(p, ', with address at ')
field(p, d.get('CLIENT_ADDRESS'), 'address'); run(p, ', ' + d['CLIENT_CITY_STATE'] + '.')

p = par(doc); run(p, 'PROVIDER: ', bold=True)
run(p, '%s, %s no. %s, with address at %s, %s.' % (d['PROVIDER_NAME'], d['PROVIDER_DOC_LABEL'], d['PROVIDER_DOC'], d['PROVIDER_ADDRESS'], d['PROVIDER_CITY_STATE']))

par(doc, 'The parties identified above enter into this service contract, which shall be governed by the following clauses.')

def clause(n, title, text):
    par(doc, 'Clause %s — %s' % (n, title), bold=True, before=12)
    par(doc, text)

clause(1, 'Purpose', 'This contract\'s purpose is the creation of a new version of the CLIENT\'s web page (%s), including: complete layout redesign preserving the visual identity (logo, colors and images provided), improved copy of the existing content, mobile adaptation and publishing of the page at %s.' % (d['OLD_SITE_URL'], d['PUBLISHED_URL']))
clause(2, 'Value and payment terms', 'For the services described in Clause 1, the CLIENT shall pay the PROVIDER the total amount of R$ %s (%s), as follows: %s.' % (d['VALUE'], d['VALUE_IN_WORDS'], d['PAYMENT_TERMS']))
clause(3, 'Delivery deadline', 'The page in its final version will be delivered and published within %s from the signing of this contract and from the CLIENT providing the necessary materials and approvals. %s round(s) of text and image adjustments after delivery are included.' % (d['DELIVERY_DEADLINE'], d['ADJUSTMENT_ROUNDS']))
n = 4
if d.get('MAINTENANCE'):
    clause(4, 'Monthly maintenance', 'The CLIENT also hires the monthly page maintenance service (hosting, small text/image updates and support), for R$ %s per month, effective from publishing with automatic monthly renewal.' % d['MAINTENANCE_VALUE'])
    n = 5
clause(n, 'Content and responsibilities', 'The CLIENT declares to be the owner or to hold authorization to use all texts, images, logo and information provided, being responsible for the truthfulness of the professional information disclosed. The PROVIDER commits not to insert on the page any information not provided or not approved by the CLIENT.')
clause(n+1, 'Hosting and domain', d['HOSTING_TEXT'])
clause(n+2, 'Termination', 'This contract may be terminated by either party through written notice. In case of termination by the CLIENT after the start of the work, the proportional value of the services already performed shall be due. Monthly maintenance services, when contracted, may be cancelled by either party with 30 (thirty) days\' prior notice.')
clause(n+3, 'Jurisdiction', 'The courts of the district of %s are chosen to settle any disputes arising from this contract.' % d['FORUM_CITY'])

p = par(doc, before=18); run(p, d['SIGNATURE_CITY'] + ', ')
editable(p, ' [date] '); run(p, '.')

par(doc, '', before=24)
p = par(doc, before=18); run(p, '__________________________________________'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = par(doc, before=0); run(p, d['CLIENT_NAME'] + ' — Client  ', bold=True); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
editable(p, ' [sign here] ')
p = par(doc, before=18); run(p, '__________________________________________'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = par(doc, before=0, center=True); run(p, d['PROVIDER_NAME'] + ' — Provider', bold=True)

par(doc, 'This document is an automatically generated base draft. Review by a legal professional before signing is recommended. Generated by the Site Prospector toolkit.', size=8, before=20)

# PROTECTION: read-only, except the permitted regions above
dp = OxmlElement('w:documentProtection')
dp.set(qn('w:edit'), 'readOnly'); dp.set(qn('w:enforcement'), '1')
doc.settings.element.append(dp)

doc.save(sys.argv[2])
print('docx generated:', sys.argv[2])
